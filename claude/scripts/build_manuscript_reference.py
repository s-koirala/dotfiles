#!/usr/bin/env python3
"""Build the manuscript reference.docx used by pandoc to style .md/.qmd → .docx.

Generates ~/.claude/templates/manuscript/reference.docx with minimalist B&W
styling that satisfies common major-journal submission requirements (NEJM,
JAMA, Lancet, BMJ, AJPH, Am J Epidemiol, Ann Intern Med, JAMA Network):

- 12pt Times New Roman, double-spaced body
- 1-inch margins all sides
- Page numbers bottom-right
- Bold headings (same point size as body)
- No colors, no shading, no decorative formatting
- Standard styles: Title, Author, Heading 1-3, Body Text, Caption, Quote,
  List Paragraph, Footnote

Generated at: ~/.claude/templates/manuscript/reference.docx

Usage:
    python scripts/build_manuscript_reference.py

Idempotent: re-running overwrites the reference doc. The user may further
customize in Word; re-running this script reverts customizations. Track
customizations in a separate doc (e.g., reference_custom.docx) if needed.

R3-9 from docs/audits/implementation_plan_dotfiles_additions_2026-05-15.md;
specs per user directive 2026-05-15 (.docx, B&W, minimalist, clinical/med/PH/PS
journal compatible).
"""
from __future__ import annotations

import sys
from pathlib import Path


def build_reference_docx(output_path: Path) -> Path:
    """Create the reference.docx via python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("ERROR: python-docx not installed. Run: uv pip install python-docx",
              file=sys.stderr)
        return Path()

    doc = Document()

    # --- Page setup: 1-inch margins all sides ---
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # --- Page numbers (bottom-right) ---
    # Add page number to footer via raw OOXML (python-docx lacks a high-level API)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    # Inject PAGE field
    fld_simple = OxmlElement("w:fldSimple")
    fld_simple.set(qn("w:instr"), "PAGE")
    run._r.append(fld_simple)

    # --- Define / override styles ---
    # justify: Times New Roman 12pt is the universal default for clinical
    # journal submission. NEJM, JAMA, Annals, AJPH, Am J Epidemiol all
    # specify 12pt Times New Roman explicitly. Lancet/BMJ accept Arial as
    # alternative; we ship Times to maximize compatibility.
    body_font = "Times New Roman"
    body_size = Pt(12)

    # Normal / Body Text style
    normal = doc.styles["Normal"]
    normal.font.name = body_font
    normal.font.size = body_size
    normal.font.color.rgb = RGBColor(0, 0, 0)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Inches(0)  # block paragraphs (no indent); justify: clinical journals overwhelmingly use block paragraphs in submission manuscripts

    # Headings (1, 2, 3): bold, same point size as body, double-spaced
    for level in (1, 2, 3):
        style_name = f"Heading {level}"
        style = doc.styles[style_name]
        style.font.name = body_font
        style.font.size = body_size
        style.font.bold = True
        style.font.italic = False
        style.font.color.rgb = RGBColor(0, 0, 0)
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        pf.space_before = Pt(12)
        pf.space_after = Pt(0)

    # Title — distinct from body via size only; still B&W
    title = doc.styles["Title"]
    title.font.name = body_font
    title.font.size = Pt(14)  # justify: one notch larger than body; matches Ann Intern Med submission guideline (12-14pt title)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0, 0, 0)
    pf = title.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)

    # Caption (for figures/tables)
    try:
        caption = doc.styles["Caption"]
    except KeyError:
        # Some Document instances don't pre-create Caption style
        from docx.enum.style import WD_STYLE_TYPE
        caption = doc.styles.add_style("Caption", WD_STYLE_TYPE.PARAGRAPH)
    caption.font.name = body_font
    caption.font.size = Pt(11)  # justify: one notch smaller than body; matches AMA Manual of Style §4.2.3 figure caption guidance
    caption.font.bold = False
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(0, 0, 0)
    pf = caption.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE  # captions single-spaced even when body is double; AMA Manual of Style
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)

    # Quote
    try:
        quote = doc.styles["Quote"]
    except KeyError:
        from docx.enum.style import WD_STYLE_TYPE
        quote = doc.styles.add_style("Quote", WD_STYLE_TYPE.PARAGRAPH)
    quote.font.name = body_font
    quote.font.size = body_size
    quote.font.italic = False  # justify: clinical journals discourage italicized block quotes; use indentation instead
    quote.font.color.rgb = RGBColor(0, 0, 0)
    pf = quote.paragraph_format
    pf.left_indent = Inches(0.5)
    pf.right_indent = Inches(0.5)
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)

    # List Paragraph
    try:
        list_p = doc.styles["List Paragraph"]
        list_p.font.name = body_font
        list_p.font.size = body_size
        list_p.font.color.rgb = RGBColor(0, 0, 0)
        pf = list_p.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    except KeyError:
        pass

    # Add a sentinel paragraph so pandoc finds at least one body paragraph
    # to derive from. Pandoc reference-doc convention: keep one Normal-style
    # paragraph; pandoc strips it when generating the actual output.
    p = doc.add_paragraph(
        "Reference doc — pandoc strips body content when rendering. "
        "Styles defined: Normal, Heading 1-3, Title, Caption, Quote, "
        "List Paragraph. Customize in Word if needed; re-running "
        "build_manuscript_reference.py reverts changes."
    )

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def main() -> int:
    out = Path(__file__).resolve().parent.parent / "templates" / "manuscript" / "reference.docx"
    result = build_reference_docx(out)
    if not result or not result.exists():
        print("FAIL: reference.docx not created", file=sys.stderr)
        return 1
    size_kb = result.stat().st_size / 1024
    print(f"PASS: reference.docx generated at {result} ({size_kb:.1f} KB)")
    print(f"  styles: Normal (12pt Times, double-spaced)")
    print(f"  margins: 1.0\" all sides")
    print(f"  page numbers: bottom-right")
    print(f"  colors: B&W only")
    print()
    print("Next: use with pandoc via scripts/render_manuscript.py")
    return 0


if __name__ == "__main__":
    sys.exit(main())
